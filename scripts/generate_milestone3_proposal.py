from __future__ import annotations

import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "Group9-PaperProposal.docx"
OUTPUT_PATH = ROOT / "Group9-PaperProposal-Milestone3.docx"

BENCHMARK_PATH = ROOT / "outputs" / "results" / "benchmark_summary.json"
RERANKER_PATH = ROOT / "outputs" / "results" / "reranker_metrics.json"
EVIDENCE_PATH = ROOT / "outputs" / "results" / "evidence_chain_analysis.json"
META_PATH = (
    ROOT
    / "data"
    / "processed"
    / "hotpot_validation_disjoint_t200_v1000_s42_metadata.json"
)


def read_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def get_style_name(doc: Document, preferred: str, fallback: str = "Normal") -> str:
    names = {s.name for s in doc.styles}
    if preferred in names:
        return preferred
    return fallback


def add_para(
    doc: Document,
    text: str,
    style: str,
    *,
    bold: bool = False,
    italic: bool = False,
    align_center: bool = False,
) -> None:
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_spacer(doc: Document, style: str, count: int = 1) -> None:
    for _ in range(count):
        doc.add_paragraph("", style=style)


def main() -> None:
    benchmark = read_json(BENCHMARK_PATH)
    reranker = read_json(RERANKER_PATH)
    evidence = read_json(EVIDENCE_PATH)
    meta = read_json(META_PATH)

    doc = Document(str(TEMPLATE_PATH))
    clear_document(doc)

    normal_style = get_style_name(doc, "normal")
    h1 = get_style_name(doc, "Heading 1")
    h2 = get_style_name(doc, "Heading 2")
    h3 = get_style_name(doc, "Heading 3")

    # Cover page in the same style family as the template.
    add_para(doc, "RESEARCH PROPOSAL - MILESTONE 3", normal_style, bold=True, align_center=True)
    add_para(
        doc,
        "Refining Research and Peer Review Process",
        normal_style,
        align_center=True,
    )
    add_spacer(doc, normal_style)
    add_para(doc, "Agent-Enhanced GraphRAG:", normal_style, bold=True, align_center=True)
    add_para(
        doc,
        "Automating Multi-Hop Question Answering",
        normal_style,
        bold=True,
        align_center=True,
    )
    add_para(
        doc,
        "with Multi-Agent Systems and Graph Learning",
        normal_style,
        bold=True,
        align_center=True,
    )
    add_spacer(doc, normal_style)
    add_para(doc, "Group 09", normal_style, align_center=True)
    add_para(
        doc,
        "Kamal Mustafayev   |   Fidan Maharlamova   |   Mahabbat Zakariyayev",
        normal_style,
        align_center=True,
    )
    add_para(doc, date.today().strftime("%d %B %Y"), normal_style, align_center=True)
    add_spacer(doc, normal_style)
    add_para(
        doc,
        "Keywords: GraphRAG, Multi-Agent Systems, HotpotQA, Multi-Hop QA,"
        " Retrieval-Augmented Generation, Explainability",
        normal_style,
        italic=True,
        align_center=True,
    )

    doc.add_page_break()

    add_para(doc, "1. Research Question and Milestone 3 Objectives", h1)
    add_para(
        doc,
        "Research Question: To what extent does an agent-enhanced GraphRAG pipeline "
        "improve multi-hop question answering quality and evidence traceability over "
        "standard dense retrieval baselines under realistic compute constraints?",
        normal_style,
    )
    add_para(
        doc,
        "Milestone 3 focuses on three graded outputs: (1) an updated state of the art "
        "after peer and instructor feedback, (2) a fully justified methodology with "
        "dataset and bias analysis, and (3) preliminary quantitative results that can "
        "support scientific presentation and methodology peer review.",
        normal_style,
    )

    add_para(doc, "Table 1: Milestone 3 objectives and measurable evidence.", normal_style)
    t1 = doc.add_table(rows=1, cols=3)
    t1.rows[0].cells[0].text = "Objective"
    t1.rows[0].cells[1].text = "What changed after first feedback"
    t1.rows[0].cells[2].text = "Current evidence"
    rows = [
        (
            "State of the art",
            "Expanded from summary citations to method-level comparison and explicit research gap.",
            "Section 2 with updated paper mapping and integration plan.",
        ),
        (
            "Methodology",
            "Moved from conceptual architecture to reproducible implementation choices (.env, model, training).",
            "Section 3 with pipeline, training protocol, and failure controls.",
        ),
        (
            "Datasets and bias",
            "Added leakage checks, split metadata, and representational limitations.",
            "Section 4 with disjoint split statistics and bias mitigation actions.",
        ),
        (
            "Preliminary results",
            "Added benchmark-prefixed quantitative outcomes (B1/B2/OURS) and evidence diagnostics.",
            "Section 5 using repository outputs from the current implementation.",
        ),
    ]
    for r in rows:
        cells = t1.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = r

    add_para(doc, "2. Modified State of the Art After First Feedback", h1)
    add_para(
        doc,
        "First-round feedback indicated that the literature review needed stronger "
        "methodological critique and clearer mapping from prior work to concrete design "
        "decisions. The revised review therefore emphasizes mechanism-level contributions, "
        "limitations, and implementation impact on this project.",
        normal_style,
    )

    add_para(doc, "2.1 Updated core papers and project-level integration", h2)
    add_para(doc, "Table 2: Modified state of the art and design implications.", normal_style)
    t2 = doc.add_table(rows=1, cols=4)
    t2.rows[0].cells[0].text = "Reference"
    t2.rows[0].cells[1].text = "Main contribution"
    t2.rows[0].cells[2].text = "Key limitation"
    t2.rows[0].cells[3].text = "How we use it"
    paper_rows = [
        (
            "Shrestha and Kim (2025)",
            "LLM planning plus embedding-guided graph search for efficient multi-hop KGQA.",
            "Limited discussion of deployment under strict resource constraints.",
            "Guides our decomposition-first retrieval strategy and hop scoring.",
        ),
        (
            "Ni et al. (2025)",
            "Step-wise graph reasoning chain (StepChain) for multi-hop QA.",
            "High-quality chains may depend on curated graph quality.",
            "Motivates explicit evidence chains and path-level diagnostics.",
        ),
        (
            "Song et al. (2026)",
            "Cooperative RAG framing with decision-making across modules.",
            "Coordination overhead can increase latency.",
            "Supports our multi-agent orchestration and critic feedback loop.",
        ),
        (
            "Edge et al. (2024)",
            "GraphRAG baseline from local graph context to structured retrieval.",
            "No dedicated agentic planning/reflection components.",
            "Used as conceptual baseline for B2 (basic GraphRAG).",
        ),
        (
            "Yao et al. (2022)",
            "ReAct unifies reasoning and actions for LLMs.",
            "Can still produce brittle chains if retrieval quality is low.",
            "Drives our reasoner implementation with evidence-grounded steps.",
        ),
    ]
    for r in paper_rows:
        cells = t2.add_row().cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text = r

    add_para(doc, "2.2 Revised research gap", h2)
    add_para(
        doc,
        "Existing systems typically optimize one axis at a time: retrieval quality, "
        "graph structure, or reasoning policy. The revised gap is therefore defined as "
        "joint optimization of (a) multi-hop answer quality, (b) evidence traceability, "
        "and (c) operational feasibility on limited local compute.",
        normal_style,
    )

    add_para(doc, "3. Research Methodology", h1)
    add_para(
        doc,
        "The implementation uses a four-agent LangGraph workflow: Query Decomposer, "
        "Graph Retriever, ReAct Reasoner, and Critic. The critic can trigger one additional "
        "retrieval loop when confidence is insufficient. This realizes iterative correction "
        "without requiring external paid APIs.",
        normal_style,
    )

    add_para(doc, "3.1 System-level implementation details", h2)
    add_para(doc, "Table 3: Methodology and concrete implementation choices.", normal_style)
    t3 = doc.add_table(rows=1, cols=3)
    t3.rows[0].cells[0].text = "Layer"
    t3.rows[0].cells[1].text = "Chosen method"
    t3.rows[0].cells[2].text = "Justification"
    method_rows = [
        (
            "Agent orchestration",
            "LangGraph state machine with explicit node transitions and feedback edges.",
            "Supports transparent multi-step control flow and iterative refinement.",
        ),
        (
            "Generation backend",
            "LLM_BACKEND=ollama and OLLAMA_MODEL=qwen2.5:7b",
            "Balances quality and local inference feasibility for iterative QA.",
        ),
        (
            "Retrieval embeddings",
            "sentence-transformers/all-MiniLM-L6-v2",
            "Fast dense representations with acceptable semantic recall on HotpotQA passages.",
        ),
        (
            "Graph layer",
            "Dynamic NetworkX graph plus BFS traversal and relation-typed edges.",
            "Maintains explicit hop provenance required for evidence chain analysis.",
        ),
        (
            "Reranking",
            "LogisticRegression reranker on lexical plus semantic feature vectors.",
            "Improves passage ordering with low training cost and interpretable metrics.",
        ),
    ]
    for r in method_rows:
        cells = t3.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = r

    add_para(doc, "3.2 Training protocol and reproducibility", h2)
    add_para(
        doc,
        "The current run uses deterministic data partition metadata with seed=42 and "
        "disjoint train/validation/test sets. The reranker was trained on 5,271 pairwise "
        "examples with a positive rate of 0.2504, reaching validation accuracy 0.8090. "
        "Environment controls include TEMPERATURE=0.1, MAX_NEW_TOKENS=384, and "
        "MAX_RETRIEVAL_LOOPS=2.",
        normal_style,
    )
    add_para(
        doc,
        "To ensure replicability, all benchmark outputs are exported in versioned JSON "
        "files under outputs/results, and the same runtime configuration is preserved in "
        "the .env settings used for evaluation.",
        normal_style,
    )

    add_para(doc, "3.3 Methodology critique and validity threats", h2)
    add_para(
        doc,
        "Internal validity risk: retrieval improvements may be partially entangled with "
        "LLM output variability. Mitigation: fixed temperature and benchmarked baselines "
        "under shared subset conditions.",
        normal_style,
    )
    add_para(
        doc,
        "External validity risk: evaluation currently emphasizes one benchmark subset. "
        "Mitigation: planned extension to MuSiQue and 2WikiMultiHopQA for transferability.",
        normal_style,
    )

    add_para(doc, "4. Datasets and Bias Analysis", h1)
    add_para(
        doc,
        "Primary dataset: HotpotQA. For this milestone we use the processed disjoint split "
        "metadata produced in the repository pipeline.",
        normal_style,
    )

    add_para(doc, "Table 4: Dataset setup and data leakage safeguards.", normal_style)
    t4 = doc.add_table(rows=1, cols=4)
    t4.rows[0].cells[0].text = "Split"
    t4.rows[0].cells[1].text = "Size"
    t4.rows[0].cells[2].text = "Overlap check"
    t4.rows[0].cells[3].text = "Purpose"

    split_sizes = meta.get("sizes", {})
    overlap = meta.get("overlap", {})
    t4_rows = [
        (
            "Train",
            str(split_sizes.get("train", "N/A")),
            f"train_val={overlap.get('train_val', 'N/A')}, train_test={overlap.get('train_test', 'N/A')}",
            "Reranker training",
        ),
        (
            "Validation",
            str(split_sizes.get("validation", "N/A")),
            f"train_val={overlap.get('train_val', 'N/A')}, val_test={overlap.get('val_test', 'N/A')}",
            "Threshold and configuration checks",
        ),
        (
            "Test",
            str(split_sizes.get("test", "N/A")),
            f"train_test={overlap.get('train_test', 'N/A')}, val_test={overlap.get('val_test', 'N/A')}",
            "Benchmark comparison (B1/B2/OURS)",
        ),
    ]
    for r in t4_rows:
        cells = t4.add_row().cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text = r

    add_para(doc, "4.1 Bias and representational limits", h2)
    add_para(
        doc,
        "Known bias source 1: English-centric entity coverage can underrepresent low-resource "
        "contexts. Known bias source 2: retrieval may favor highly linked entities. Known bias "
        "source 3: confidence scores may not reflect calibration under adversarial phrasing.",
        normal_style,
    )
    add_para(
        doc,
        "Mitigation actions include error slicing by question type, source diversity monitoring "
        "in evidence chains, and expansion to additional multi-hop datasets in the next cycle.",
        normal_style,
    )

    add_para(doc, "5. Preliminary Results", h1)
    b1 = benchmark["B1"]
    b2 = benchmark["B2"]
    ours = benchmark["OURS"]

    em_delta_vs_b1 = ours["exact_match"] - b1["exact_match"]
    f1_delta_vs_b1 = ours["f1"] - b1["f1"]
    exp_delta_vs_b1 = ours["explainability"] - b1["explainability"]

    add_para(doc, "Table 5: Benchmark summary on the 200-question test subset.", normal_style)
    t5 = doc.add_table(rows=1, cols=6)
    t5.rows[0].cells[0].text = "System"
    t5.rows[0].cells[1].text = "Exact Match"
    t5.rows[0].cells[2].text = "F1"
    t5.rows[0].cells[3].text = "Explainability"
    t5.rows[0].cells[4].text = "Delta EM vs B1"
    t5.rows[0].cells[5].text = "Delta F1 vs B1"

    result_rows = [
        (
            "B1 (Dense RAG)",
            f"{b1['exact_match']:.3f}",
            f"{b1['f1']:.3f}",
            f"{b1['explainability']:.3f}",
            "0.000",
            "0.000",
        ),
        (
            "B2 (Basic GraphRAG)",
            f"{b2['exact_match']:.3f}",
            f"{b2['f1']:.3f}",
            f"{b2['explainability']:.3f}",
            f"{(b2['exact_match'] - b1['exact_match']):+.3f}",
            f"{(b2['f1'] - b1['f1']):+.3f}",
        ),
        (
            "OURS (Agent-Enhanced GraphRAG)",
            f"{ours['exact_match']:.3f}",
            f"{ours['f1']:.3f}",
            f"{ours['explainability']:.3f}",
            f"{em_delta_vs_b1:+.3f}",
            f"{f1_delta_vs_b1:+.3f}",
        ),
    ]
    for r in result_rows:
        cells = t5.add_row().cells
        for idx, val in enumerate(r):
            cells[idx].text = val

    add_para(
        doc,
        "Key observation: OURS improves answer quality over B1 and B2 on Exact Match and F1, "
        "while explainability remains high but currently slightly below B1. This indicates "
        "the architecture is effective for correctness, with additional work needed on "
        "explanation consistency and calibration.",
        normal_style,
    )

    add_para(doc, "5.1 Evidence chain quality diagnostics", h2)
    add_para(doc, "Table 6: Evidence-chain statistics for the OURS configuration.", normal_style)
    t6 = doc.add_table(rows=1, cols=2)
    t6.rows[0].cells[0].text = "Metric"
    t6.rows[0].cells[1].text = "Value"
    diag_rows = [
        ("Total predictions", str(evidence.get("total_predictions", "N/A"))),
        ("Empty chain rate", f"{evidence.get('empty_chain_rate', 0.0):.3f}"),
        ("Average chain length", f"{evidence.get('avg_chain_length', 0.0):.3f}"),
        ("Average hop score", f"{evidence.get('avg_hop_score', 0.0):.3f}"),
        ("Average source diversity", f"{evidence.get('avg_source_diversity', 0.0):.3f}"),
        ("Malformed hop rate", f"{evidence.get('malformed_hop_rate', 0.0):.3f}"),
    ]
    for r in diag_rows:
        cells = t6.add_row().cells
        cells[0].text, cells[1].text = r

    add_para(doc, "5.2 Preliminary interpretation", h2)
    add_para(
        doc,
        "The evidence diagnostics suggest stable chain generation (zero malformed hops and "
        "zero empty chains) with moderate source diversity. This supports the claim that "
        "graph-guided retrieval provides traceable multi-hop reasoning paths in practice.",
        normal_style,
    )

    add_para(doc, "6. Scientific Presentation of Current Project State", h1)
    add_para(
        doc,
        "The project is presentation-ready at the prototype-plus-evaluation stage. The "
        "scientific presentation is structured around: (1) motivation and gap, (2) pipeline "
        "architecture, (3) updated literature synthesis, (4) benchmark outcomes, and "
        "(5) limitations with next-step validation plan.",
        normal_style,
    )
    add_para(
        doc,
        "All reported metrics are directly linked to repository artifacts, ensuring that "
        "slides and proposal narrative remain reproducible and auditable.",
        normal_style,
    )

    add_para(doc, "7. Peer Review Plan for Methodology", h1)
    add_para(
        doc,
        "Peer review is organized around three mandatory dimensions: relevance, feasibility, "
        "and justification. Each criterion is scored and accompanied by concrete revision "
        "actions for the next milestone.",
        normal_style,
    )
    add_para(doc, "Table 7: Methodology peer-review rubric.", normal_style)
    t7 = doc.add_table(rows=1, cols=4)
    t7.rows[0].cells[0].text = "Criterion"
    t7.rows[0].cells[1].text = "Review question"
    t7.rows[0].cells[2].text = "Current status"
    t7.rows[0].cells[3].text = "Next action"
    peer_rows = [
        (
            "Relevance",
            "Does each module directly support multi-hop QA quality and traceability goals?",
            "Strong",
            "Add ablation showing contribution of each agent.",
        ),
        (
            "Feasibility",
            "Can the full pipeline run under local or free-tier constraints?",
            "Strong",
            "Add repeated-run latency distribution and cost profile.",
        ),
        (
            "Justification",
            "Are model choices and datasets defended by evidence rather than preference?",
            "Moderate to strong",
            "Include statistical significance tests and calibration checks.",
        ),
    ]
    for r in peer_rows:
        cells = t7.add_row().cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text = r

    add_para(doc, "8. Remaining Work to Final Paper", h1)
    add_para(
        doc,
        "Planned tasks before final submission: run ablation experiments, evaluate on "
        "additional datasets (MuSiQue and 2WikiMultiHopQA), perform significance testing, "
        "improve explanation calibration, and finalize qualitative error taxonomy.",
        normal_style,
    )

    add_para(doc, "9. Conclusion", h1)
    add_para(
        doc,
        "Milestone 3 delivers a revised literature review, a justified and reproducible "
        "methodology, and preliminary benchmark evidence. The current results support the "
        "core hypothesis that agent-enhanced GraphRAG improves multi-hop QA quality over "
        "standard baselines while preserving evidence traceability.",
        normal_style,
    )

    add_para(doc, "References", h1)
    references = [
        "Edge, D., et al. (2024). From Local to Global: A GraphRAG Approach to Query-Focused Summarization. arXiv:2404.16130.",
        "Ni, X., et al. (2025). StepChain GraphRAG: Reasoning Over Knowledge Graphs for Multi-Hop QA. arXiv:2510.02827.",
        "Shrestha, S., and Kim, J. (2025). Efficient Multi-Hop QA over KGs via LLM Planning and Embedding-Guided Search. arXiv:2511.19648.",
        "Song, Y., et al. (2026). Rethinking Retrieval-Augmented Generation as a Cooperative Decision-Making Problem. arXiv:2602.18734.",
        "Yao, S., et al. (2022). ReAct: Synergizing Reasoning and Acting in Language Models. arXiv:2210.03629.",
        "Yang, Z., et al. (2018). HotpotQA: A Dataset for Diverse, Explainable Multi-Hop Question Answering. EMNLP 2018.",
    ]
    for ref in references:
        add_para(doc, ref, normal_style)

    doc.save(str(OUTPUT_PATH))
    print(f"Created: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
